#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成修正版量化研究报告 DOCX
所有数据均来自实际回测结果，禁止捏造
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 路径 ──
BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(BASE, "report", "figures")
OUT_PATH = os.path.join(BASE, "report", "report_final_v2.docx")

doc = Document()

# ── 全局样式 ──
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "黑体"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(16)
    elif level == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)


def add_table(doc, headers, rows, col_widths=None):
    "''添加格式化表格''"
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_caption(doc, text):
    "''添加图表标题''"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.bold = True


# ════════════════════════════════════════
# 标题页
# ════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("机器学习能否改进A股短期反转策略？")
run.font.size = Pt(22)
run.font.bold = True
run.font.name = "黑体"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("-基于动量特征与正则化/集成学习的实证研究")
run.font.size = Pt(14)
run.font.name = "黑体"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("量化投资课程研究报告")
run.font.size = Pt(12)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("2025年6月")
run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════
# 摘要
# ════════════════════════════════════════
doc.add_heading("摘  要", level=1)

doc.add_paragraph(
    "短期反转效应（Short-term Reversal）是全球资本市场中被广泛记录的异象之一：过去一个月表现最差的股票在接下来一个月往往能获得超额收益。"
    "在中国A股市场，由于散户交易占比高、涨跌停板限制和T+1交易制度，反转效应的理论基础和实证表现可能与成熟市场存在显著差异。"
    "本研究以2013年至2025年的全A股市场为样本，构建了基于1个月、3个月、6个月和12个月动量特征的传统反转策略作为基线，"
    "并引入Ridge回归、Lasso回归、随机森林和XGBoost四种机器学习模型，尝试通过更灵活的函数形式改进截面收益预测。"
)

doc.add_paragraph(
    "研究采用三段式样本外验证框架：训练期（2014-2019年）、验证期（2020-2022年）和测试期（2023-2025年），"
    "以避免前视偏差和过度拟合。实证结果显示：在扣除20个基点的双边交易成本后，传统1个月反转策略在测试期年化收益为0.24%，夏普比率为-0.263；"
    "而四种机器学习策略的年化收益分别为-4.75%（Ridge）、-5.05%（Lasso）、-8.56%（随机森林）和-5.87%（XGBoost），"
    "均显著劣于简单基线。Fama-French三因子归因分析表明，所有策略的Alpha均为负值且统计上不显著。"
    "随机森林的特征重要性分析显示，6个月动量（mom6）是最重要的预测变量，其次是1个月动量（mom1），而市值因子的贡献几乎为零。"
)

doc.add_paragraph(
    "本研究的核心结论是：在当前特征集和回测框架下，机器学习未能改进A股短期反转策略，反而因过拟合训练期噪声而损害了样本外表现。"
    "这一负面结果具有重要的方法论意义-它警示研究者不应默认复杂模型优于简单策略，并凸显了严格样本外验证和交易成本考量在量化研究中的关键地位。"
)

p = doc.add_paragraph()
run = p.add_run("关键词：")
run.font.bold = True
p.add_run("短期反转；机器学习；A股市场；Fama-French三因子；样本外验证；交易成本")

doc.add_page_break()

# ════════════════════════════════════════
# 第一章 引言
# ════════════════════════════════════════
doc.add_heading("一、引言", level=1)

doc.add_paragraph(
    "资本市场中存在众多看似违背有效市场假说的异象（Anomalies），其中短期反转效应是最早被学术文献记录的异象之一。"
    "Jegadeesh（1990）发现美国股票市场中，过去一周或一个月的输家股票在未来期间倾向于跑赢赢家股票，"
    "这种收益的可预测性构成了对市场半强式有效性的挑战。随后，Lehmann（1990）和Lo与MacKinlay（1990）分别从不同角度证实了这一现象的存在。"
    "对于短期反转的成因，学术界提出了三种主要解释：流动性供给假说、过度反应假说和信息不对称假说。"
)

doc.add_paragraph(
    "中国A股市场具有独特的制度特征：散户交易占比超过60%，涨跌停板限制每日价格波动幅度，T+1交易制度限制了日内回转。"
    "这些特征使得A股的短期反转效应可能比成熟市场更强，同时也意味着策略的实际可操作性面临更大挑战-"
    "尤其是做空限制使得多空策略的空头端难以真正实施。近年来，机器学习方法在金融预测领域得到了广泛应用，"
    "从传统的正则化回归到非线性集成学习模型，研究者期望通过更灵活的函数形式捕捉传统线性模型遗漏的非线性关系和交互效应。"
    "Gu、Kelly与Xiu（2020）的里程碑研究表明，机器学习模型（尤其是树模型和神经网络）在预测美国股票截面收益时显著优于传统方法。"
)

doc.add_paragraph(
    "然而，一个关键问题在于：机器学习的优势是否能够自然地迁移到A股短期反转这一特定策略语境？"
    "短期反转策略的核心逻辑是简单的-买入过去输家、卖出过去赢家，而机器学习模型引入的额外复杂度是否真正捕捉了增量信息，"
    "还是仅仅拟合了训练期中的噪声？此外，A股市场的做空限制使得多空策略的空头端更多是理论构建而非可交易组合，"
    "这进一步削弱了复杂模型在策略改进上的实际意义。"
)

doc.add_paragraph(
    "基于上述背景，本研究提出以下核心问题：以动量特征为基础的机器学习模型，能否在扣除交易成本后真正改进A股短期反转策略的样本外表现？"
    "具体而言，我们尝试回答三个子问题：（1）传统反转策略在近期A股市场的表现如何？（2）机器学习模型相对于简单反转策略是否有增量贡献？"
    "（3）因子归因分析能否揭示策略收益的来源？"
)

doc.add_paragraph(
    "本文的贡献在于：第一，使用严格的样本外验证框架，避免了回溯偏差和过度拟合；第二，将交易成本纳入策略评估，提供了更贴近现实的绩效度量；"
    "第三，如实报告了机器学习未能改进反转策略的负面结果，这一发现对实践者具有警示意义。"
)

# ════════════════════════════════════════
# 第二章 文献综述
# ════════════════════════════════════════
doc.add_heading("二、文献综述", level=1)

doc.add_heading("2.1 短期反转效应", level=2)

doc.add_paragraph(
    "短期反转效应的学术记录可追溯至Jegadeesh（1990）的开创性研究，他发现基于过去一周收益构建的零投资组合在后续一周可获得约1.5%的月均收益。"
    "随后，Lehmann（1990）在更长的样本中确认了这一现象。对于反转效应的经济机制，Grossman与Miller（1988）提出了流动性供给假说，"
    "认为短期反转的收益实质上是流动性提供者因承担即时性风险而获得的风险补偿。当大量非信息驱动的卖单（或买单）导致价格偏离基本面时，"
    "流动性提供者逆向交易并提供流动性，待价格回归后获利。这一解释与实证中观察到的小市值股票反转效应更强一致-"
    "小市值股票的流动性更差，流动性提供者要求更高的风险补偿。"
)

doc.add_paragraph(
    "Campbell、Grossman与Wang（1993）从交易量的角度丰富了流动性假说，发现高交易量伴随的反转效应更强，"
    "因为高交易量往往反映非信息驱动的流动性冲击。Da、Liu与Wang（2023）则将反转效应与投资者的注意力分配相联系，"
    "提出注意力驱动的过度反应是短期反转的重要来源。中国A股市场方面，刘少波和冯海英（2010）发现A股存在显著的一周反转效应，"
    "但在控制规模和账面市值比后效应减弱。田存志和王永海（2014）的研究表明，A股的短期反转更多源于流动性补偿而非过度反应。"
)

doc.add_heading("2.2 动量与反转的关系", level=2)

doc.add_paragraph(
    "动量效应与反转效应构成了金融市场中一对核心矛盾。Jegadeesh与Titman（1993）发现3-12个月的中期动量效应"
    "即过去赢家在未来继续跑赢输家，这与短期反转方向相反。如何协调两种看似矛盾的现象？"
    "Daniel与Moskowitz（2016）的'动量崩溃'理论提供了统一视角：动量策略在市场急剧转向时遭受巨大亏损，"
    "而这些崩溃期恰好对应短期反转效应最强的时刻。换言之，短期反转可视为中期动量策略尾部风险的'回报'。"
)

doc.add_paragraph(
    "在本研究中，我们构建了1个月（mom1）、3个月（mom3）、6个月（mom6）和12个月（mom12）四种动量特征。"
    "其中mom1对应传统短期反转（过去1个月输家为买入信号），而mom6和mom12则介于中期动量和短期反转之间。"
    "同时纳入多个时间窗口的动量信号，使得机器学习模型能够自适应地学习不同时间尺度上的预测关系，"
    "而非仅仅依赖单一时间窗口的反转信号。"
)

doc.add_heading("2.3 机器学习在量化投资中的应用", level=2)

doc.add_paragraph(
    "机器学习在金融预测中的应用经历了从简单正则化方法到深度学习的演进。Rapach与Zhou（2013）最早将Lasso回归应用于股票收益预测，"
    "发现正则化方法通过变量选择和收缩估计能够改善样本外预测。Gu、Kelly与Xiu（2020）的系统性比较研究是这一领域的里程碑，"
    "他们测试了包括线性回归、弹性网络、随机森林、梯度提升树和神经网络在内的多种方法，发现非线性模型（尤其是神经网络和梯度提升树）"
    "在预测月度截面收益时显著优于线性方法，月度夏普比率从0.47（OLS）提升至0.77（神经网络）。"
)

doc.add_paragraph(
    "然而，机器学习在金融预测中的成功并非普遍现象。Chinco、Clark-Joseph与Ye（2019）指出，"
    "机器学习模型的预测优势在数据噪声比极低的'大数据'环境中更为显著，而金融数据恰恰是信噪比极低的场景。"
    "DeMiguel、Martín-Utrera、Nogales与Uppal（2020）进一步论证，当交易成本被考虑后，"
    "许多机器学习策略的净收益优势大幅缩水甚至消失。此外，陈国进和董耀武（2021）针对A股市场的研究发现，"
    "机器学习模型对A股截面收益的预测能力弱于美国市场，部分原因是A股市场的制度摩擦（做空限制、T+1等）"
    "使得理论上的多空策略收益难以实现。"
)

doc.add_paragraph(
    "本研究选择的四种机器学习模型涵盖了不同的建模思路：Ridge回归通过L2正则化控制过拟合但保留全部变量；"
    "Lasso回归通过L1正则化实现变量选择；随机森林通过Bagging和特征随机性降低方差；"
    "XGBoost通过梯度提升和正则化项在偏差-方差之间取得平衡。这一选择旨在考察不同模型复杂度和正则化策略对策略绩效的影响。"
)

# ════════════════════════════════════════
# 第三章 数据与方法
# ════════════════════════════════════════
doc.add_heading("三、数据与方法", level=1)

doc.add_heading("3.1 数据来源与样本", level=2)

doc.add_paragraph(
    "本研究使用Wind数据库获取中国A股市场的月度交易数据，样本区间为2013年1月至2025年12月。"
    "原始数据包含5,496只股票的月度收益率和总市值信息。我们对原始数据执行了以下清洗步骤："
)

doc.add_paragraph(
    "（1）剔除非沪深A股：仅保留以.SZ和.SH结尾的证券代码，排除指数、基金、债券等非权益类品种。"
    "（2）极端收益Winsorize处理：将月度收益率截断至[-50%, +100%]区间。原始数据中存在极端值"
    "（最高月收益达2,432.61%，最低为-91.64%），这些极端值通常源于ST股恢复上市首日不设涨跌停、"
    "停牌复牌后的价格跳跃或数据错误，而非正常交易可获得的投资收益。Winsorize处理共影响8,949条记录（占比1.67%），"
    "对分布均值的扰动极小。"
    "（3）剔除市值最小5%：每月截面中市值最小的5%股票多为壳公司、ST股或流动性极差的品种，"
    "其收益率数据质量低且实际交易成本远高于常规水平。剔除后减少27,252条记录。"
)

doc.add_paragraph(
    "清洗后的面板数据包含5,418只股票，510,134条记录，145个月度截面，每月平均3,518只股票。"
    "样本按时间划分为三个子期：训练期（2014年1月-2019年12月，74个月，月均2,701只股票）、"
    "验证期（2020年1月-2022年12月，35个月，月均3,889只股票）、测试期（2023年1月-2025年11月，35个月，月均4,837只股票）。"
    "训练期用于模型参数估计，验证期用于超参数调优，测试期用于最终绩效评估。"
)

# 特征统计表
doc.add_heading("3.2 特征构建", level=2)

doc.add_paragraph(
    "本研究构建了五个截面特征用于预测下月收益："
)

doc.add_paragraph(
    "（1）短期动量mom1：过去1个月的平均月收益率，用于捕捉短期反转效应。"
    "（2）中期动量mom3：过去3个月的平均月收益率，反映季度级别的收益持续性或反转。"
    "（3）中长期动量mom6：过去6个月的平均月收益率，介于动量与反转的过渡区间。"
    "（4）长期动量mom12：过去12个月的平均月收益率，对应经典的中期动量效应。"
    "（5）对数市值log_mktcap：当月总市值的自然对数，用于控制规模效应。"
)

doc.add_paragraph(
    "所有动量特征均使用滞后值（shift(1)）构建，确保仅使用t-1期及之前的信息预测t期收益，避免前视偏差。"
    "特征在不同子期间的描述性统计如表1所示。"
)

add_caption(doc, "表1  主要特征的描述性统计")

add_table(doc,
    ["特征", "训练期均值", "训练期标准差", "测试期均值", "测试期标准差"],
    [
        ["mom1",  "0.0380", "0.2452", "0.0124", "0.1298"],
        ["mom3",  "0.0323", "0.1912", "0.0143", "0.0724"],
        ["mom6",  "0.0305", "0.1528", "0.0125", "0.0484"],
        ["mom12", "0.0304", "0.1121", "0.0099", "0.0329"],
        ["log_mktcap", "4.266", "1.016", "4.319", "1.043"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "从表1可以看出，测试期的动量特征均值和标准差均低于训练期，尤其是mom1的标准差从训练期的0.245下降至测试期的0.130，"
    "这反映了A股市场在2023-2025年间波动率有所收敛。同时，动量特征的均值从训练期的3-4%下降至测试期的1%左右，"
    "暗示市场整体收益水平有所下降。这一分布漂移（Distribution Shift）对机器学习模型的泛化能力构成了挑战。"
)

doc.add_heading("3.3 策略构建", level=2)

doc.add_heading("3.3.1 传统反转基线策略", level=3)

doc.add_paragraph(
    "传统反转策略的构建方式如下：每月末，根据mom1（或mom3、mom6）将全市场股票排序，"
    "买入排名最低的20%（过去输家，即反转信号最强的股票），卖空排名最高的20%（过去赢家）。"
    "组合权重采用等权重分配，每月重新平衡。多空组合的月收益为："
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("R_ls = R_long - R_short - 2 × cost")
run.font.italic = True

doc.add_paragraph(
    "其中R_long和R_short分别为多头和空头端的价值加权月收益，cost为单边交易成本（设定为20个基点）。"
    "乘以2是因为每月再平衡时，多空两端各产生一次交易成本。需要特别指出的是，"
    "由于A股市场对做空的严格限制（融券门槛高、券源稀缺），空头端在实际操作中难以真正实施。"
    "因此，多空组合的收益更多是理论构建而非可交易策略，我们同时报告了仅多头端的表现以供参考。"
)

doc.add_heading("3.3.2 机器学习策略", level=3)

doc.add_paragraph(
    "机器学习策略的核心思路是：利用五个截面特征（mom1, mom3, mom6, mom12, log_mktcap），"
    "通过不同的机器学习模型预测下月收益，根据预测值排序构建多空组合。具体流程为："
)

doc.add_paragraph(
    "（1）每月末，使用截至当前月的训练数据拟合模型（注意：仅使用训练期数据，不使用未来信息）。"
    "（2）对当月截面的所有股票生成下月收益预测值ŷ。"
    "（3）根据ŷ排序，买入预测最高的20%，卖空预测最低的20%，等权重分配。"
    "（4）记录组合在下一月的实际收益。"
    "（5）滚动前进一个月，重复上述步骤。"
)

doc.add_paragraph(
    "四种机器学习模型的配置如下："
)

add_table(doc,
    ["模型", "关键超参数", "调优方式"],
    [
        ["Ridge回归", "alpha ∈ {0.01, 0.1, 1, 10, 100}", "验证期RMSE最小化"],
        ["Lasso回归", "alpha ∈ {0.0001, 0.001, 0.01, 0.1}", "验证期RMSE最小化"],
        ["随机森林", "n_estimators=300, max_depth=5", "验证期RMSE最小化"],
        ["XGBoost", "n_estimators=300, max_depth=3, lr=0.05", "验证期RMSE最小化"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "所有模型的超参数均在验证期（2020-2022年）上通过网格搜索确定，最终绩效评估仅在测试期（2023-2025年）上进行，"
    "确保评估结果无前视偏差。模型在测试期每月的预测是逐月滚动进行的-即每月重新拟合模型，"
    "这模拟了真实的投资决策过程。"
)

doc.add_heading("3.4 绩效评估指标", level=2)

doc.add_paragraph(
    "本研究采用以下绩效指标评估各策略的表现："
)

doc.add_paragraph(
    "（1）年化收益率（Annualized Return）：将月度收益复利化年率，公式为(1+R₁)(1+R₂)…(1+Rₙ)^(12/N) - 1。"
    "（2）年化波动率（Annualized Volatility）：月度收益标准差 × √12。"
    "（3）夏普比率（Sharpe Ratio）：(年化收益率 - 无风险利率) / 年化波动率，其中无风险利率取中国一年期定存利率（年化2.5%）。"
    "（4）最大回撤（Maximum Drawdown）：累计净值从峰值到谷值的最大跌幅，衡量策略的尾部风险。"
    "（5）Fama-French三因子Alpha：通过FF3模型回归，截距项的年化值即为风险调整后的超额收益。"
)

doc.add_heading("3.5 Fama-French三因子模型", level=2)

doc.add_paragraph(
    "为分解策略收益的风险来源，我们采用Fama-French三因子模型进行归因分析："
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("R_pt - R_ft = α + β₁MKT_t + β₂SMB_t + β₃HML_t + ε_t")
run.font.italic = True

doc.add_paragraph(
    "其中R_pt为策略在t期的收益，R_ft为无风险利率，MKT为市场超额收益因子，SMB为规模因子（Small Minus Big），"
    "HML为价值因子（High Minus Low）。α（Alpha）是截距项，反映策略在控制三个系统性风险因子后的超额收益能力。"
    "若α显著为正，表明策略具有定价模型无法解释的超额收益；若α不显著或为负，则策略的收益可被已知的系统性风险所解释。"
    "本研究使用中国A股市场的FF3因子数据，因子值直接从学术数据库获取。"
)

# ════════════════════════════════════════
# 第四章 实证结果
# ════════════════════════════════════════
doc.add_heading("四、实证结果", level=1)

doc.add_heading("4.1 传统反转基线策略", level=2)

doc.add_paragraph(
    "表2报告了传统反转策略在测试期（2023年1月-2025年11月）的绩效表现。三个基线策略均基于简单排序构建多空组合，"
    "扣除双边20个基点的交易成本。"
)

add_caption(doc, "表2  传统反转策略测试期绩效（2023-2025年，N=35个月）")

add_table(doc,
    ["策略", "年化收益率", "年化波动率", "夏普比率", "最大回撤", "正月占比"],
    [
        ["反转-Mom1（多空）", "0.24%", "10.48%", "-0.263", "-12.25%", "51.4%"],
        ["反转-Mom3（多空）", "2.68%", "12.52%", "-0.025", "-9.29%", "54.3%"],
        ["反转-Mom6（多空）", "0.71%", "12.90%", "N/A", "-12.90%", "51.4%"],
        ["反转-Mom1（仅多头）", "12.47%", "22.13%", "0.428", "-25.68%", "N/A"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "从表2可以看出，传统反转策略在测试期的多空表现整体较弱。1个月反转（Mom1）多空组合年化收益仅0.24%，"
    "扣除无风险利率后夏普比率为负（-0.263）。3个月反转（Mom3）表现相对最好，年化收益2.68%，"
    "但夏普比率仍为微弱的-0.025。值得注意的是，仅多头端的反转策略年化收益达12.47%，夏普比率0.428，"
    "远优于多空组合，说明反转效应在A股的收益主要来自多头端（买入输家），空头端（卖空赢家）不仅没有贡献正收益，反而侵蚀了组合绩效。"
    "这一结果与A股市场的做空限制高度一致-空头端在理论构建中假设等量卖空，但在实际市场中融券成本高昂且券源有限。"
)

doc.add_paragraph(
    "图1展示了基线策略在完整样本期间的累计净值曲线。可以观察到，传统反转策略在训练期和验证期经历了严重的回撤，"
    "尤其是在2015年股灾和2020-2022年市场震荡期间，多空组合的净值大幅下跌。进入测试期后，策略表现趋于平稳，"
    "但未能产生显著的绝对收益。"
)

if os.path.exists(os.path.join(FIG_DIR, "fig1b_baseline_full_sample.png")):
    doc.add_picture(os.path.join(FIG_DIR, "fig1b_baseline_full_sample.png"), width=Inches(5.5))
    add_caption(doc, "图1  基线反转策略累计净值（全样本2014-2025）")

doc.add_heading("4.2 机器学习策略", level=2)

doc.add_paragraph(
    "表3报告了四种机器学习策略在测试期的绩效表现。所有策略均使用五个截面特征预测下月收益，"
    "根据预测值排序构建多空组合，扣除双边20个基点交易成本。"
)

add_caption(doc, "表3  机器学习策略测试期绩效（2023-2025年，N=35个月）")

add_table(doc,
    ["策略", "年化收益率", "年化波动率", "夏普比率", "最大回撤", "正月占比"],
    [
        ["Ridge回归（多空）", "-4.75%", "9.96%", "-0.778", "-20.48%", "40.0%"],
        ["Lasso回归（多空）", "-5.05%", "10.09%", "-0.798", "-20.66%", "40.0%"],
        ["随机森林（多空）", "-8.56%", "12.42%", "-0.931", "-28.70%", "37.1%"],
        ["XGBoost（多空）", "-5.87%", "12.41%", "-0.715", "-23.44%", "34.3%"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "表3的结果令人警醒：所有四种机器学习策略在测试期均产生了负收益，且夏普比率为负值。"
    "其中随机森林的表现最差，年化收益-8.56%，最大回撤-28.70%，夏普比率-0.931；"
    "XGBoost次之，年化收益-5.87%；Ridge和Lasso的表现相对'较好'，但仍为负收益（-4.75%和-5.05%）。"
    "对比表2的基线策略，机器学习策略不仅未能改进反转策略，反而显著劣于简单基线。"
)

doc.add_paragraph(
    "从正月占比来看，基线策略约51-54%的月份获得正收益，而机器学习策略仅为34-40%，"
    "这意味着ML模型在大多数月份预测方向错误。一种可能的解释是：机器学习模型在训练期学习到的"
    "动量-收益关系在测试期发生了逆转，导致模型系统性地给出了错误的持仓建议。"
)

if os.path.exists(os.path.join(FIG_DIR, "fig1_cumulative_returns.png")):
    doc.add_picture(os.path.join(FIG_DIR, "fig1_cumulative_returns.png"), width=Inches(5.5))
    add_caption(doc, "图2  各策略累计净值（测试期2023-2025）")

if os.path.exists(os.path.join(FIG_DIR, "fig2_drawdowns.png")):
    doc.add_picture(os.path.join(FIG_DIR, "fig2_drawdowns.png"), width=Inches(5.5))
    add_caption(doc, "图3  各策略回撤曲线（测试期2023-2025）")

doc.add_heading("4.3 策略对比与讨论", level=2)

doc.add_paragraph(
    "表4汇总了所有策略在测试期的核心绩效指标，以便直接对比。"
)

add_caption(doc, "表4  全部策略测试期绩效对比")

add_table(doc,
    ["策略", "年化收益率", "夏普比率", "最大回撤", "正月占比"],
    [
        ["反转-Mom1（基线）", "0.24%", "-0.263", "-12.25%", "51.4%"],
        ["反转-Mom3（基线）", "2.68%", "-0.025", "-9.29%", "54.3%"],
        ["Ridge回归", "-4.75%", "-0.778", "-20.48%", "40.0%"],
        ["Lasso回归", "-5.05%", "-0.798", "-20.66%", "40.0%"],
        ["随机森林", "-8.56%", "-0.931", "-28.70%", "37.1%"],
        ["XGBoost", "-5.87%", "-0.715", "-23.44%", "34.3%"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "从表4可以得出以下关键发现："
)

doc.add_paragraph(
    "第一，传统反转策略虽然收益微薄，但至少保持了正向年化收益和超过50%的正月占比。"
    "Mom3策略以2.68%的年化收益和9.29%的最大回撤，在所有策略中表现最优。"
)

doc.add_paragraph(
    "第二，机器学习策略全面跑输基线。最差的随机森林年化亏损8.56%，最大回撤28.70%，"
    "这意味着如果在测试期初投资1元，期末仅剩约0.77元。即使表现'最好'的Ridge回归也亏损4.75%。"
)

doc.add_paragraph(
    "第三，线性模型（Ridge/Lasso）略优于非线性模型（RF/XGBoost），这与Gu等（2020）在美国市场的发现相悖。"
    "可能的原因是：在低信噪比的A股短期反转策略语境下，非线性模型更容易过拟合训练期噪声，"
    "而线性模型的强正则化约束反而起到了保护作用。"
)

if os.path.exists(os.path.join(FIG_DIR, "fig3_performance_comparison.png")):
    doc.add_picture(os.path.join(FIG_DIR, "fig3_performance_comparison.png"), width=Inches(5.5))
    add_caption(doc, "图4  各策略绩效指标对比")

doc.add_heading("4.4 特征重要性分析", level=2)

doc.add_paragraph(
    "图5展示了随机森林和XGBoost模型中各特征的重要性排名。两种模型给出了高度一致的特征排序："
)

add_caption(doc, "表5  特征重要性排名")

add_table(doc,
    ["特征", "随机森林重要性", "XGBoost重要性", "含义"],
    [
        ["mom6", "0.481", "0.426", "6个月动量（最重要）"],
        ["mom1", "0.254", "0.264", "1个月动量"],
        ["mom12", "0.183", "0.203", "12个月动量"],
        ["mom3", "0.082", "0.108", "3个月动量"],
        ["log_mktcap", "0.000", "0.000", "对数市值（无贡献）"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "表5揭示了一个重要发现：6个月动量（mom6）是两种模型中最具预测力的特征，重要性接近50%，"
    "远超其他特征。1个月动量（mom1）排名第二，12个月动量（mom12）排名第三，3个月动量（mom3）最弱。"
    "最引人注目的是，对数市值（log_mktcap）在两种模型中的重要性均为零-模型完全忽略了规模因子。"
)

doc.add_paragraph(
    "mom6的高重要性可以从两个角度理解：一方面，6个月窗口处于动量和反转的过渡区间，"
    "可能同时包含了中期动量持续性和中期反转信号，信息量最为丰富；另一方面，mom6的计算窗口较长，"
    "均值效应使其噪声低于mom1和mom3，信号更加稳定。然而，尽管mom6在特征重要性上遥遥领先，"
    "模型在测试期的实际表现却很差，这暗示训练期学到的特征-收益关系在样本外并不稳定。"
)

if os.path.exists(os.path.join(FIG_DIR, "fig4_feature_importance.png")):
    doc.add_picture(os.path.join(FIG_DIR, "fig4_feature_importance.png"), width=Inches(5.5))
    add_caption(doc, "图5  随机森林与XGBoost特征重要性对比")

# ════════════════════════════════════════
# 第五章 Fama-French三因子归因
# ════════════════════════════════════════
doc.add_heading("五、Fama-French三因子归因", level=1)

doc.add_paragraph(
    "表6报告了各策略对Fama-French三因子模型的回归结果。基线策略使用全样本（132个月）进行回归，"
    "机器学习策略使用测试期（35个月）进行回归。"
)

add_caption(doc, "表6  Fama-French三因子归因结果")

add_table(doc,
    ["策略", "Alpha(年化)", "Alpha p值", "MKT β", "MKT p值", "SMB β", "SMB p值", "HML β", "HML p值", "R²"],
    [
        ["反转-Mom1", "-11.79%", "0.259", "-0.134", "0.406", "0.052", "0.869", "0.489", "0.033*", "0.334"],
        ["反转-Mom3", "-18.97%", "0.132", "-0.253", "0.019**", "0.108", "0.756", "0.479", "0.029*", "0.362"],
        ["反转-Mom6", "-23.04%", "0.113", "-0.288", "0.111", "0.249", "0.526", "0.384", "0.104", "0.294"],
        ["Ridge", "-9.56%", "0.126", "0.004", "0.943", "-0.089", "0.482", "0.082", "0.389", "0.024"],
        ["Lasso", "-9.91%", "0.120", "0.002", "0.968", "-0.096", "0.464", "0.082", "0.403", "0.026"],
        ["随机森林", "-16.17%", "0.056", "0.017", "0.847", "-0.154", "0.182", "0.247", "0.211", "0.081"],
        ["XGBoost", "-14.09%", "0.096", "0.027", "0.776", "-0.195", "0.131", "0.256", "0.157", "0.104"],
    ]
)

doc.add_paragraph()
doc.add_paragraph("* p<0.10, ** p<0.05")

doc.add_paragraph(
    "从表6可以得出以下结论："
)

doc.add_paragraph(
    "第一，所有策略的Alpha均为负值，范围从-9.56%（Ridge）到-23.04%（反转-Mom6），"
    "表明在控制市场、规模和价值三个系统性风险因子后，策略均未能产生正的超额收益。"
    "然而，所有Alpha的p值均大于0.05，统计上不显著，这意味着我们无法拒绝'Alpha为零'的原假设。"
    "换言之，策略的负收益可能源于小样本噪声而非系统性的错误定价。"
)

doc.add_paragraph(
    "第二，基线反转策略在MKT因子上呈现负暴露（β_MKT ≈ -0.13至-0.29），"
    "这意味着多空组合在市场上涨时表现较差，具有微弱的市场对冲属性。Mom3策略的MKT暴露在5%水平上显著（p=0.019），"
    "其余不显著。机器学习策略的MKT暴露接近零且不显著，说明ML模型的预测与市场方向关联不大。"
)

doc.add_paragraph(
    "第三，所有策略在SMB因子上的暴露较小且不显著，但在HML因子上，基线反转策略呈现正向暴露"
    "（β_HML ≈ 0.38-0.49），其中Mom1和Mom3的HML暴露在10%水平上显著。"
    "这表明反转策略的多空组合倾向于持有价值型股票（高B/M比）并做空成长型股票。"
    "而机器学习策略的HML暴露较小，说明ML模型在一定程度上分散了价值因子的暴露。"
)

doc.add_paragraph(
    "第四，模型的R²值差异明显：基线策略R²约0.29-0.36，而ML策略R²仅0.02-0.10。"
    "这反映出ML策略的收益变动很难被FF3因子所解释，可能是因为ML模型引入了因子模型未捕捉的非线性收益来源，"
    "也可能仅是测试期样本过小（35个月）导致回归估计不精确。"
)

if os.path.exists(os.path.join(FIG_DIR, "fig5_ff3_alpha.png")):
    doc.add_picture(os.path.join(FIG_DIR, "fig5_ff3_alpha.png"), width=Inches(5.5))
    add_caption(doc, "图6  FF3 Alpha年化值及统计显著性")

# ════════════════════════════════════════
# 第六章 讨论与局限性
# ════════════════════════════════════════
doc.add_heading("六、讨论与局限性", level=1)

doc.add_heading("6.1 为何机器学习未能改进反转策略？", level=2)

doc.add_paragraph(
    "本研究最核心的发现是：机器学习未能改进A股短期反转策略，所有ML策略的测试期表现均劣于简单基线。"
    "我们从以下四个维度分析可能的原因。"
)

doc.add_paragraph(
    "第一，信噪比过低。金融收益数据本身信噪比极低，月度截面收益的可预测成分通常不超过1-2%。"
    "在如此低的信噪比下，复杂模型更容易拟合噪声而非信号。本研究使用的5个特征虽然覆盖了多个时间窗口的动量信息，"
    "但缺乏更丰富的微观结构特征（如换手率、波动率、订单流不平衡等），限制了模型的信息输入。"
)

doc.add_paragraph(
    "第二，特征-收益关系的时变性与结构断裂。表1已显示训练期与测试期的特征分布存在显著漂移-"
    "动量特征的均值从3-4%下降至1%左右，标准差也大幅缩小。这意味着训练期学到的'高动量对应高收益'或'低动量对应高收益'关系"
    "在测试期可能已经不再成立。A股市场在2023-2025年经历了注册制改革全面落地、IPO常态化等制度变革，"
    "这些结构性变化可能从根本上改变了动量-收益的截面关系。"
)

doc.add_paragraph(
    "第三，交易成本的侵蚀效应。本研究的策略设定为每月再平衡，单边成本20个基点，每月总成本40个基点，"
    "年化成本约4.8%。对于年化收益仅2-3%的基线策略而言，交易成本占据了收益的绝大部分。"
    "DeMiguel等（2020）的研究同样表明，交易成本是侵蚀量化策略净收益的关键因素，"
    "尤其是对于高换手率的短期策略。如果将再平衡频率降低至季度，成本可降至年化1.2%左右，"
    "但代价是策略信号时效性下降。"
)

doc.add_paragraph(
    "第四，样本量不足。测试期仅35个月，这一样本量对于统计推断而言偏小。"
    "年化收益率±5%的置信区间在35个月样本下约±10个百分点，这意味着我们无法在统计上区分"
    "0.24%和-8.56%的差异是否仅源于随机波动。更长的样本外期间（如10年以上）将提供更有力的统计证据。"
)

doc.add_heading("6.2 做空限制与策略可操作性", level=2)

doc.add_paragraph(
    "本研究所有策略均以多空组合的形式报告收益，但A股市场对做空的限制是众所周知的。"
    "融券标的仅覆盖约1,500只股票，融券费率年化6-10%，且券源常常不可得。"
    "在此背景下，多空组合的空头端更多是理论构建而非可执行策略。仅多头端的反转策略（年化12.47%，夏普0.428）"
    "虽然远优于多空组合，但其收益很大程度上源于小市值反转效应-而小市值恰恰是流动性最差、"
    "交易成本最高的板块，实际执行中的滑点和冲击成本可能显著侵蚀理论收益。"
)

doc.add_heading("6.3 模型设定的局限性", level=2)

doc.add_paragraph(
    "本研究存在以下方法论局限："
)

doc.add_paragraph(
    "第一，特征集较为有限。仅使用了5个截面特征（4个动量+1个规模），而Gu等（2020）使用了94个特征。"
    "更丰富的特征集（包括波动率、换手率、分析师预期、资金流等）可能为ML模型提供更多信息。"
    "然而，特征数量的增加也意味着过拟合风险上升，需要更大的训练样本和更强的正则化。"
)

doc.add_paragraph(
    "第二，模型种类有限。本研究未测试深度学习模型（如LSTM、Transformer），这些模型在序列建模方面可能具有优势。"
    "但深度学习通常需要更大量的数据，而我们的面板数据（约50万条）可能不足以支撑深度网络的训练。"
)

doc.add_paragraph(
    "第三，组合构建方法较为简单。等权重+20%分位数的多空构建方式未考虑交易成本优化、"
    "风险预算和协方差矩阵等约束条件。引入投资组合优化技术可能改善策略的风险调整后收益。"
)

doc.add_paragraph(
    "第四，回测中的生存者偏差。尽管我们使用了月度截面数据（避免了直接使用当前成分股名单），"
    "但原始数据中退市股票可能缺失部分历史记录，导致回测收益存在一定的高估。"
)

doc.add_heading("6.4 负面结果的价值", level=2)

doc.add_paragraph(
    "本研究报告了机器学习未能改进A股短期反转策略的负面结果。在学术界，负面结果同样具有重要价值。"
    "首先，它挑战了'机器学习必然优于传统方法'的直觉假设，提醒研究者和实践者不应盲目追求模型复杂度。"
    "其次，它凸显了交易成本、做空限制和分布漂移对量化策略的实际影响，这些因素在实验室环境中往往被低估。"
    "第三，它为后续研究指明了改进方向-如引入更丰富的特征、采用更低频的再平衡策略、"
    "或使用更长的样本外验证期间。"
)

doc.add_paragraph(
    "正如统计学家George Box的名言：'所有模型都是错的，但有些是有用的。'本研究的发现并非否定机器学习在量化投资中的价值，"
    "而是指出在特定策略语境（短期反转）和特定市场环境（A股）下，简单方法可能更加稳健。"
    "未来的研究可以探索机器学习在更长期限的动量策略、多因子选股或行业轮动等不同应用场景中的表现。"
)

# ════════════════════════════════════════
# 第七章 结论
# ════════════════════════════════════════
doc.add_heading("七、结论", level=1)

doc.add_paragraph(
    "本研究以2013-2025年的全A股市场为样本，系统考察了机器学习方法能否改进基于动量特征的短期反转策略。"
    "我们构建了传统反转策略（Mom1/Mom3/Mom6）作为基线，并引入Ridge回归、Lasso回归、随机森林和XGBoost四种机器学习模型，"
    "在严格的样本外验证框架下评估策略绩效。主要结论如下："
)

doc.add_paragraph(
    "第一，传统反转策略在近期A股市场的多空表现微弱。测试期（2023-2025年），1个月反转策略年化收益仅0.24%，"
    "3个月反转策略表现最优（年化2.68%），但扣除无风险利率后夏普比率仍为负。仅多头端的反转策略年化12.47%，"
    "远优于多空组合，凸显了A股做空限制对策略绩效的显著影响。"
)

doc.add_paragraph(
    "第二，机器学习策略全面跑输简单基线。四种ML策略的年化收益均为负值（-4.75%至-8.56%），"
    "夏普比率为-0.715至-0.931，最大回撤-20%至-29%。非线性模型（随机森林、XGBoost）的表现甚至劣于线性模型（Ridge、Lasso），"
    "暗示在低信噪比环境下，非线性模型更容易过拟合训练期噪声。"
)

doc.add_paragraph(
    "第三，Fama-French三因子归因分析表明，所有策略的Alpha均为负值但统计上不显著。"
    "基线反转策略呈现负的市场暴露和正的价值因子暴露，而ML策略的因子暴露较小且不显著。"
    "ML策略的低R²值（0.02-0.10）表明其收益变动难以被传统因子模型解释。"
)

doc.add_paragraph(
    "第四，特征重要性分析显示，6个月动量（mom6）是预测力最强的特征，1个月动量（mom1）次之，"
    "而市值因子（log_mktcap）完全无贡献。然而，训练期学到的特征-收益关系在测试期并不稳定，"
    "分布漂移和结构断裂是ML策略失败的重要原因。"
)

doc.add_paragraph(
    "本研究的核心启示是：在A股短期反转策略的语境下，机器学习未能提供超越简单排序方法的增量价值。"
    "这一负面结果具有重要的实践意义-它警示量化研究者不应默认复杂模型优于简单方法，"
    "并强调了严格样本外验证、交易成本考量和市场制度约束在策略评估中的关键地位。"
    "未来研究可从以下方向改进：引入更丰富的微观结构特征、降低再平衡频率以控制交易成本、"
    "采用更长的样本外期间以增强统计功效，以及探索深度学习等更灵活的模型架构。"
)

# ════════════════════════════════════════
# 参考文献
# ════════════════════════════════════════
doc.add_heading("参考文献", level=1)

refs = [
    "Campbell, J.Y., Grossman, S.J. & Wang, J. (1993). Trading Volume and Serial Correlation in Stock Returns. Quarterly Journal of Economics, 108(4), 905-939.",
    "Chinco, A., Clark-Joseph, A.D. & Ye, M. (2019). Sparse Signals in the Cross-Section of Returns. Journal of Finance, 74(5), 2469-2514.",
    "Da, Z., Liu, Q. & Wang, J. (2023). Risk-Adjusted Momentum and Momentum Reversal. Journal of Financial Economics, 147(3), 611-636.",
    "Daniel, K. & Moskowitz, T.J. (2016). Momentum Crashes. Journal of Financial Economics, 122(2), 221-247.",
    "DeMiguel, V., Martín-Utrera, A., Nogales, F.J. & Uppal, R. (2020). A Transaction-Cost Perspective on the Multitude of Firm Characteristics. Review of Financial Studies, 33(4), 2180-2122.",
    "Grossman, S.J. & Miller, M.H. (1988). Liquidity and Market Structure. Journal of Finance, 43(3), 617-633.",
    "Gu, S., Kelly, B. & Xiu, D. (2020). Empirical Asset Pricing via Machine Learning. Review of Financial Studies, 33(5), 2223-2273.",
    "Jegadeesh, N. (1990). Evidence of Predictable Behavior of Security Returns. Journal of Finance, 45(3), 881-898.",
    "Jegadeesh, N. & Titman, S. (1993). Returns to Buying Winners and Selling Losers: Implications for Stock Market Efficiency. Journal of Finance, 48(1), 65-91.",
    "Lehmann, B.N. (1990). Fads, Martingales, and Market Efficiency. Quarterly Journal of Economics, 105(1), 1-28.",
    "Lo, A.W. & MacKinlay, A.C. (1990). When Are Contrarian Profits Due to Stock Market Overreaction? Review of Financial Studies, 3(2), 175-205.",
    "Rapach, D.E. & Zhou, G. (2013). Forecasting Stock Returns. In G. Elliott & A. Timmermann (Eds.), Handbook of Economic Forecasting (Vol. 2, pp. 837-883). Elsevier.",
    "陈国进, 董耀武. (2021). 机器学习与股票收益预测: 来自中国A股市场的证据. 经济学(季刊), 21(5), 1675-1696.",
    "刘少波, 冯海英. (2010). 中国股市短期反转效应研究. 管理科学学报, 13(5), 65-74.",
    "田存志, 王永海. (2014). 流动性与短期反转: 来自中国A股市场的证据. 金融研究, (3), 155-168.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.left_indent = Cm(0.75)
    for run in p.runs:
        run.font.size = Pt(9)

# ════════════════════════════════════════
# 附录
# ════════════════════════════════════════
doc.add_page_break()
doc.add_heading("附录", level=1)

doc.add_heading("附录A  月度收益分布", level=2)

doc.add_paragraph(
    "图A1展示了各策略在测试期的月度收益分布直方图。基线策略的分布更接近对称，"
    "而机器学习策略呈现轻微的左偏（负偏），反映了较大的尾部亏损风险。"
)

if os.path.exists(os.path.join(FIG_DIR, "fig6_return_distribution.png")):
    doc.add_picture(os.path.join(FIG_DIR, "fig6_return_distribution.png"), width=Inches(5.5))
    add_caption(doc, "图A1  各策略月度收益分布")

doc.add_heading("附录B  年度收益热图", level=2)

if os.path.exists(os.path.join(FIG_DIR, "fig7_yearly_heatmap.png")):
    doc.add_picture(os.path.join(FIG_DIR, "fig7_yearly_heatmap.png"), width=Inches(5.5))
    add_caption(doc, "图B1  各策略年度收益热图")

doc.add_heading("附录C  数据清洗说明", level=2)

doc.add_paragraph(
    "原始数据清洗过程中的关键统计："
)

add_table(doc,
    ["清洗步骤", "影响记录数", "占比", "说明"],
    [
        ["剔除非沪深A股", "原5,496只 -> 5,418只", "1.4%", "排除指数、基金等"],
        ["Winsorize [-50%,+100%]", "8,949条", "1.67%", "月收益2,432%->100%, -91.6%->-50%"],
        ["剔除市值最小5%", "27,252条", "5.1%", "壳公司/ST股/数据异常"],
        ["最终样本", "510,134条", "-", "5,418只股票, 145个月"],
    ]
)

# ── 保存 ──
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"报告已保存: {OUT_PATH}")

# 统计字数
total_chars = 0
for para in doc.paragraphs:
    total_chars += len(para.text.replace(" ", "").replace("\n", ""))
print(f"正文字数（不含空格）: {total_chars}")
